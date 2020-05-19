import requests
from bs4 import BeautifulSoup
import xlrd
import xlwt
import numpy as np
from scipy.optimize import curve_fit
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
#pip install xlrd
#pip install xlwt
#pip install bs4
#本爬虫抓取数据源自Wikipedia，如需复现，请自行代理，或直接使用境外服务器执行
#本爬虫在阿里云香港ecs debian10 python3.7.3环境下测试通过

def func1(x, a, b,c):
    return a*x+b+c
def func2(x, a, b,c):
    return a*x*x+b*x*x+c
def func3(x, a, b,c):
    return a*x*x*x+b*x*x+c
def __main__():
    date = []
    people = []
    #发送get请求并使用bs4进行处理
    res = requests.get('https://zh.wikipedia.org/wiki/%E4%B8%AD%E5%9B%BD%E5%A4%A7%E9%99%86%E4%BA%BA%E5%8F%A3')
    res.encoding = 'utf-8'
    soup=BeautifulSoup(res.text,'html.parser')
    #查找中国人口变化的table
    table = soup.find('table',class_ = 'wikitable collapsible')
    #print(table)
    temp = table.get_text()
    #对table中的text进行预处理
    temp = ''.join(temp).split()
    for i in range(len(temp)):
        try:
            int(temp[i])
        except:
            continue
        if int(temp[i]) < 2020 and int(temp[i]) > 1950:
            ttt = list(temp[i+1])
            for j in range(len(ttt)):
                if ttt[j] == ',':
                    ttt=ttt[:j]+ttt[j+1:]
                    break
            temp[i+1] = int(''.join(ttt))
            date.append(temp[i])
            people.append(temp[i+1])
    print(date)
    print(people)
    workbook = xlwt.Workbook()  # 新建一个工作簿
    sheet = workbook.add_sheet('China')  # 在工作簿中新建一个表格
    for i in range(len(date)):
        sheet.write(i,0,date[i])
        sheet.write(i,1,people[i])
    workbook.save('data.xls')  # 保存工作簿

    #以下进行数据处理
    data = xlrd.open_workbook("data.xls")
    table = data.sheets()[0]
    
    
    #从excel表格中读取数据
    x = table.col_values(0)
    y = table.col_values(1)
    for i in range(len(x)):
        x[i] = float(x[i])
        y[i] = float(y[i])
    y = np.array(y)
    x = np.array(x)
    
    

    #获取popt里面是拟合系数
    #分别使用一次、二次、三次函数拟合
    popt, pcov = curve_fit(func1, x, y)
    print(popt)
    a = popt[0] 
    b = popt[1]
    c = popt[2]
    yvals = func1(x,a,b,c)
    plot2 = plt.plot(x, yvals, 'r',label='linear function')

    popt, pcov = curve_fit(func2, x, y)
    print(popt)
    a = popt[0] 
    b = popt[1]
    c = popt[2]
    yvals = func2(x,a,b,c)
    plot2 = plt.plot(x, yvals, 'b',label='Quadratic function')
    popt, pcov = curve_fit(func3, x, y)
    print(popt)
    a = popt[0] 
    b = popt[1]
    c = popt[2]
    yvals = func3(x,a,b,c)
    plot2 = plt.plot(x, yvals, 'g',label='Three functions')
    #作图
    plot1 = plt.plot(x, y, 's',label='original values')
    plt.xlabel('x')
    plt.ylabel('y')
    plt.legend(loc=4) #指定legend的位置右下角
    plt.title('''A Scatter Chart of China's Population Over Time''')
    plt.savefig("data.jpg")
    
    
    #将图片从jpg转存到Word
    doc = Document()    # doc对象
    doc.add_paragraph('中国人口随时间变化曲线图')   # 添加文字
    doc.add_picture('data.jpg', width=Inches(5))     # 添加图, 设置宽度
    doc.save('data.docx')
if __name__ == "__main__":
    __main__()